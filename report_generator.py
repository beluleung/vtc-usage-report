import io
import os
import argparse
from datetime import datetime, timedelta, UTC
from typing import Dict, List, Optional, Tuple
from decimal import Decimal

import boto3
import pandas as pd
from botocore.config import Config
from dotenv import load_dotenv

# Optional dependency imports for DOCX export. Imported lazily in exporter.


# ---------------------------
# Configuration
# ---------------------------
METRICS_USAGE_TYPE_MAP: Dict[str, List[str]] = {
    "Logins": [],
    "Uploads": [],
    "Generated Transcripts": ["transcript"],
    "Regenerated Transcripts": ["regenerate transcript"],
    "Initial Summaries": ["initial summary"],
    "Regenerated Summaries": ["regenerate summary"],
    "Regenerated Notes": ["regenerate note"],
}

TABLE_ACCOUNTS = "oak-account-vtc"
TABLE_USAGE = "oak-usage-log-vtc"
TABLE_ASKAI = "oak-ask-ai-vtc"

ACCOUNT_ALIASES: List[str] = [
    "account", "email", "user", "user_id", "userId",
    "user_email", "userEmail", "emailAddress", "user_email_address",
]
USERNAME_ALIASES: List[str] = [
    "username", "user_name", "name", "displayName",
]
USAGE_TYPE_ALIASES: List[str] = [
    "usage_type", "usageType", "type", "action", "event", "event_type",
]
CREATED_AT_ALIASES: List[str] = [
    "createdAt", "created_at", "timestamp", "createdOn",
    "created_at_ms", "created_at_iso",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate VTC OAK usage report (Excel or DOCX) from DynamoDB tables",
    )
    parser.add_argument(
        "--start-date", type=str, help="Start date in YYYY-MM-DD (defaults to 30 days ago)",
    )
    parser.add_argument(
        "--end-date", type=str, help="End date in YYYY-MM-DD (defaults to today)",
    )
    parser.add_argument(
        "--format", choices=["excel", "docx"], default="excel", help="Output format (excel or docx)",
    )
    parser.add_argument(
        "--output", type=str, default=None, help="Output filename (e.g., vtc_report.xlsx or vtc_report.docx)",
    )
    parser.add_argument(
        "--debug", action="store_true", help="Print debug info about detected columns and row counts",
    )
    parser.add_argument(
        "--no-date-filter", action="store_true", help="Disable date filtering (use all rows) for debugging",
    )
    return parser.parse_args()


def resolve_date_range(start_date_str: Optional[str], end_date_str: Optional[str]) -> Tuple[pd.Timestamp, pd.Timestamp]:
    """Resolve input date strings to inclusive [start, end] UTC timestamps."""
    if end_date_str:
        end_dt = pd.to_datetime(end_date_str, utc=True) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)
    else:
        # Fixed to use modern, non-deprecated datetime method
        end_dt = pd.to_datetime(datetime.now(UTC), utc=True).normalize() + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)

    if start_date_str:
        start_dt = pd.to_datetime(start_date_str, utc=True)
    else:
        start_dt = (end_dt - pd.Timedelta(days=29)).normalize()

    return start_dt, end_dt


def load_env_credentials() -> Tuple[str, str, str]:
    load_dotenv()
    access_key = os.getenv("VTC_DYNAMODB_ACCESS_KEY_ID")
    secret_key = os.getenv("VTC_DYNAMODB_SECRET_ACCESS_KEY")
    region = os.getenv("VTC_DYNAMODB_REGION")
    if not access_key or not secret_key or not region:
        raise RuntimeError(
            "Missing AWS credentials in .env file. Ensure VTC_DYNAMODB_... variables are set."
        )
    return access_key, secret_key, region


def get_data_from_dynamodb(table_name: str, access_key: str, secret_key: str, region: str) -> pd.DataFrame:
    """Scan a DynamoDB table and return results as a pandas DataFrame."""
    session = boto3.session.Session(
        aws_access_key_id=access_key, aws_secret_access_key=secret_key, region_name=region
    )
    dynamodb = session.resource("dynamodb", config=Config(retries={"max_attempts": 10, "mode": "standard"}))
    table = dynamodb.Table(table_name)
    items: List[dict] = []
    scan_kwargs: Dict[str, object] = {}
    while True:
        response = table.scan(**scan_kwargs)
        items.extend(response.get("Items", []))
        last_key = response.get("LastEvaluatedKey")
        if not last_key:
            break
        scan_kwargs["ExclusiveStartKey"] = last_key
    if not items:
        return pd.DataFrame()
    return pd.DataFrame(items)


def _find_first_column(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    return None


# [FIXED] This function is now robust and correctly handles the Unix timestamps
# stored as 'object' or 'Decimal' types from DynamoDB.
def coerce_created_at(df: pd.DataFrame) -> pd.DataFrame:
    """Converts various timestamp formats into a timezone-aware UTC datetime column."""
    if df.empty:
        return df
    created_col = _find_first_column(df, CREATED_AT_ALIASES)
    if not created_col:
        return df

    df = df.copy()

    # First, attempt to convert the column to a numeric type. This will handle
    # Decimals, numbers stored as strings, etc. 'coerce' turns failures into NaT.
    numeric_timestamps = pd.to_numeric(df[created_col], errors='coerce')

    # Check if the conversion to numeric was successful for most of the data.
    if numeric_timestamps.notna().mean() > 0.5:
        # Data is numeric (Unix timestamp). Now determine if it's seconds or milliseconds.
        # Heuristic: if the max value is larger than a plausible seconds value for the next 50 years,
        # assume it's milliseconds. 1e12 is a safe threshold.
        if numeric_timestamps.max() > 1e12:
             coerced = pd.to_datetime(numeric_timestamps, unit='ms', utc=True, errors='coerce')
        else:
             coerced = pd.to_datetime(numeric_timestamps, unit='s', utc=True, errors='coerce')
    else:
        # If the data is not numeric, it must be string-based (e.g., ISO 8601 format).
        coerced = pd.to_datetime(df[created_col], utc=True, errors='coerce')

    df["createdAt_dt"] = coerced
    return df


def filter_by_date(df: pd.DataFrame, start_dt: pd.Timestamp, end_dt: pd.Timestamp) -> pd.DataFrame:
    """Filters a DataFrame between a start and end timezone-aware timestamp."""
    if df.empty:
        return df
    if "createdAt_dt" not in df.columns:
        df = coerce_created_at(df)
    if "createdAt_dt" not in df.columns or df["createdAt_dt"].isna().all():
        # If column is still missing or all values are invalid, return empty.
        return df.iloc[0:0]
    mask = (df["createdAt_dt"] >= start_dt) & (df["createdAt_dt"] <= end_dt)
    return df.loc[mask].copy()


def count_usage_by_account(usage_df: pd.DataFrame, usage_types: List[str]) -> pd.DataFrame:
    if usage_df.empty or not usage_types:
        return pd.DataFrame(columns=["account", "count"])
    account_col = _find_first_column(usage_df, ACCOUNT_ALIASES)
    usage_type_col = _find_first_column(usage_df, USAGE_TYPE_ALIASES)
    if not account_col or not usage_type_col:
        return pd.DataFrame(columns=["account", "count"])
    df_norm = usage_df.copy()
    df_norm[account_col] = df_norm[account_col].astype(str).str.strip().str.lower()
    subset = df_norm[df_norm[usage_type_col].isin(usage_types)]
    grouped = subset.groupby(account_col).size().reset_index(name="count")
    grouped = grouped.rename(columns={account_col: "account"})
    return grouped


def count_askai_by_account(askai_df: pd.DataFrame) -> pd.DataFrame:
    if askai_df.empty:
        return pd.DataFrame(columns=["account", "count"])
    account_col = _find_first_column(askai_df, ACCOUNT_ALIASES)
    if not account_col:
        return pd.DataFrame(columns=["account", "count"])
    df_norm = askai_df.copy()
    df_norm[account_col] = df_norm[account_col].astype(str).str.strip().str.lower()
    grouped = df_norm.groupby(account_col).size().reset_index(name="count")
    grouped = grouped.rename(columns={account_col: "account"})
    return grouped


def build_report_dataframe(
    accounts_df: pd.DataFrame, usage_df: pd.DataFrame, askai_df: pd.DataFrame
) -> pd.DataFrame:
    if accounts_df.empty:
        columns = [
            "Account", "Username", "Logins", "Uploads", "Generated Transcripts", "Regenerated Transcripts",
            "Initial Summaries", "Regenerated Summaries", "Regenerated Notes", "AskAI Questions"
        ]
        return pd.DataFrame(columns=columns)
    base = accounts_df.copy()
    acct_col = _find_first_column(base, ACCOUNT_ALIASES)
    user_col = _find_first_column(base, USERNAME_ALIASES)
    if acct_col and acct_col != "account":
        base = base.rename(columns={acct_col: "account"})
    if user_col and user_col != "username":
        base = base.rename(columns={user_col: "username"})
    if "account" not in base.columns:
        base["account"] = None
    if "username" not in base.columns:
        base["username"] = None
    base["account"] = base["account"].astype(str).str.strip().str.lower()
    base = base[~base["account"].str.endswith("@thinkcol.com", na=False)]
    report = base[["account", "username"]].rename(columns={"account": "Account", "username": "Username"}).copy()
    for metric_name, usage_types in METRICS_USAGE_TYPE_MAP.items():
        grouped = count_usage_by_account(usage_df, usage_types)
        grouped = grouped.rename(columns={"account": "Account", "count": metric_name})
        report = report.merge(grouped, on="Account", how="left")
    askai_grouped = count_askai_by_account(askai_df).rename(columns={"account": "Account", "count": "AskAI Questions"})
    report = report.merge(askai_grouped, on="Account", how="left")
    columns_to_fill = [
        "Logins", "Uploads", "Generated Transcripts", "Regenerated Transcripts", "Initial Summaries",
        "Regenerated Summaries", "Regenerated Notes", "AskAI Questions"
    ]
    for col in columns_to_fill:
        if col not in report.columns:
            report[col] = 0
        report[col] = report[col].fillna(0).astype(int)
    if "Username" in report.columns:
        report = report.sort_values(by=["Username", "Account"], kind="stable", na_position="last").reset_index(drop=True)
    else:
        report = report.sort_values(by=["Account"], kind="stable").reset_index(drop=True)
    report = report[[
        "Account", "Username", "Logins", "Uploads", "Generated Transcripts", "Regenerated Transcripts",
        "Initial Summaries", "Regenerated Summaries", "Regenerated Notes", "AskAI Questions",
    ]]
    return report


def export_excel(df: pd.DataFrame) -> bytes:
    """Exports a DataFrame to an in-memory Excel file (bytes)."""
    output = io.BytesIO()
    # Use the 'xlsxwriter' engine for better compatibility
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='VTC Usage Report')
    # Get the content of the BytesIO object
    data = output.getvalue()
    return data


def export_docx(df: pd.DataFrame, start_dt: pd.Timestamp, end_dt: pd.Timestamp) -> bytes:
    """Exports a DataFrame to an in-memory DOCX file (bytes)."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output = io.BytesIO()
    doc = Document()

    # --- Logo ---
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "vtc_logo.png")
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Title & Subtitle ---
    title = doc.add_paragraph()
    run_t = title.add_run("VTC OAK Usage Report")
    run_t.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    range_text = f"{start_dt.strftime('%Y-%m-%d')} to {end_dt.strftime('%Y-%m-%d')} (UTC)"
    ts_text = datetime.now(UTC).strftime("%Y-%m-%d %H:%M:%S UTC")
    subtitle.add_run(f"Date Range: {range_text}    Generated: {ts_text}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("") # Spacer

    # --- Table ---
    columns = list(df.columns)
    table = doc.add_table(rows=len(df) + 1, cols=len(columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(columns):
        hdr_run = hdr_cells[idx].paragraphs[0].add_run(str(col))
        hdr_run.bold = True

    for r_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for c_idx, col in enumerate(columns):
            table.rows[r_idx].cells[c_idx].text = str(row[col]) if pd.notna(row[col]) else ""

    # Save the document to the in-memory stream
    doc.save(output)
    data = output.getvalue()
    return data


def main() -> None:
    args = parse_args()
    start_dt, end_dt = resolve_date_range(args.start_date, args.end_date)
    if args.format == "excel":
        default_name = f"vtc_report_{start_dt.strftime('%Y%m%d')}_{end_dt.strftime('%Y%m%d')}.xlsx"
    else:
        default_name = f"vtc_report_{start_dt.strftime('%Y%m%d')}_{end_dt.strftime('%Y%m%d')}.docx"
    output_path = args.output or default_name

    access_key, secret_key, region = load_env_credentials()

    accounts_df = get_data_from_dynamodb(TABLE_ACCOUNTS, access_key, secret_key, region)
    usage_df_all = get_data_from_dynamodb(TABLE_USAGE, access_key, secret_key, region)
    askai_df_all = get_data_from_dynamodb(TABLE_ASKAI, access_key, secret_key, region)

    if args.no_date_filter:
        usage_df = coerce_created_at(usage_df_all)
        askai_df = coerce_created_at(askai_df_all)
    else:
        usage_df = filter_by_date(usage_df_all, start_dt, end_dt)
        askai_df = filter_by_date(askai_df_all, start_dt, end_dt)

    if args.debug:
        print("--- DEBUG INFO ---")
        print(f"Date Range: {start_dt} to {end_dt}")
        print(f"Accounts Fetched: {len(accounts_df)}")
        print(f"Total Usage Rows Fetched: {len(usage_df_all)}")
        print(f"Usage Rows After Filtering: {len(usage_df)}")
        print(f"Total AskAI Rows Fetched: {len(askai_df_all)}")
        print(f"AskAI Rows After Filtering: {len(askai_df)}")
        if not usage_df.empty:
            ut_col = _find_first_column(usage_df, USAGE_TYPE_ALIASES)
            if ut_col:
                print(f"\nValue counts of '{ut_col}' in FILTERED data:")
                print(usage_df[ut_col].value_counts())
        print("--- END DEBUG INFO ---")


    report_df = build_report_dataframe(accounts_df, usage_df, askai_df)

    if args.format == "excel":
        export_excel(report_df, output_path)
        print(f"✅ Excel report written to {output_path}")
    else:
        export_docx(report_df, output_path, start_dt, end_dt)
        print(f"✅ DOCX report written to {output_path}")


if __name__ == "__main__":
    main()
